"""

Explains how to read and write .docx files using python
how to install - pip install python-docx
Documentation: python-docx.readthedocs.org

"""
import docx

d=docx.Document("filename") # gives a list of paragraph objects 

d.paragraphs[0].text  # gives the text in that paragraph

p =d.paragraph[1] # gives the runs associated with the paragraph

p.runs[0].text # gives the text associated with the line

p.runs[0].bold # gives True if the text in run is bold
p.runs[0].undelrine
p.runs[0].italic

p.runs[0].text="italic and underlined" # changes text to italic with underline
p.style = "Title"   # changes to the text to the Title
d.save("file_name")


# writing data to the document

d = docx.Document()
d.add_paragraph("text you want to write")
d.add_paragraph("text you want to write")
p = d.paragraphs[0]  # add text to the first paragraph
p.add_run('This is a new run')
p.runs
p.runs[1].bold=True

d.save("file_name")


# Getting all the data in the document as  a single string

import docx

def getText(filename):
    doc=docx.Document("file_name")
    fullText=[]
    for paragraph in doc.paragraphs:
      fullText.append(paragraph.text)
    return '\n'.join(fullText)
    
 print(getText(file_name))
    
